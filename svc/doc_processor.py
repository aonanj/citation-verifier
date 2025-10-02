# Copyright © 2025 Phaethon Order LLC. All rights reserved. Provided solely for evaluation. See LICENSE.

import os
import re
from collections import Counter
from typing import Any, Dict, Final, Iterable, List, Optional, Sequence, Set
from xml.etree import ElementTree as ET

import fitz  # PyMuPDF
import pytesseract
from docx import Document
from docx.document import Document as DocxDocument
from docx.oxml.ns import qn
from docx.oxml.table import CT_Tbl  # type: ignore
from docx.oxml.text.paragraph import CT_P  # type: ignore
from docx.table import Table, _Cell  # type: ignore
from docx.text.paragraph import Paragraph
from PIL import Image
from werkzeug.datastructures import FileStorage

from utils.logger import get_logger

logger = get_logger()

_HYPHEN_WRAP_RE: Final = re.compile(r"(\w)-\n(\w)")
_EXCESS_BREAKS_RE: Final = re.compile(r"\n{3,}")
_SMART_QUOTES_RE: Final = re.compile("[\u201c\u201d]")
_SMART_APOSTROPHES_RE: Final = re.compile("[\u2018\u2019]")
_SUPERSCRIPT_TRANSLATION: Final = str.maketrans("⁰¹²³⁴⁵⁶⁷⁸⁹", "0123456789")
_SUPERSCRIPT_CHARACTERS: Final = frozenset("⁰¹²³⁴⁵⁶⁷⁸⁹")
_FOOTNOTE_LINE_RE: Final = re.compile(r"^\s*([\d⁰¹²³⁴⁵⁶⁷⁸⁹]+)[\.\)]?\s*(.*)")

_W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
_NS = {"w": _W_NS}


def _load_footnotes_map(docx: DocxDocument) -> Dict[int, str]:
    """Parse word/footnotes.xml and return {footnote_id: text}."""
    mapping: Dict[int, str] = {}
    package_part = getattr(docx, "part", None)
    if package_part is None:
        return mapping
    package = getattr(package_part, "package", None)
    if package is None:
        return mapping
    footnotes_part = None
    for part in package.iter_parts():
        if str(part.partname) == "/word/footnotes.xml":
            footnotes_part = part
            break
    if footnotes_part is None:
        return mapping
    root = ET.fromstring(footnotes_part.blob)
    for fn in root.findall("w:footnote", _NS):
        fid = int(fn.get(f"{{{_W_NS}}}id", "-1"))
        if fid < 0:
            continue  # skip separators/continuation
        # Collect text paragraph-by-paragraph to preserve basic structure
        paras: List[str] = []
        for p in fn.findall(".//w:p", _NS):
            runs = [t.text or "" for t in p.findall(".//w:t", _NS)]
            txt = "".join(runs).strip()
            if txt:
                paras.append(txt)
        if not paras:
            # Fallback: any text nodes
            paras = [t.text or "" for t in fn.findall(".//w:t", _NS)]
        mapping[fid] = "\n".join([t for t in paras if t]).strip()
    return mapping

def _iter_table_paragraphs(tbl: Table) -> Iterable[Paragraph]:
    for row in tbl.rows:
        for cell in row.cells:
            yield from _iter_cell_paragraphs(cell)


def _iter_cell_paragraphs(cell: _Cell) -> Iterable[Paragraph]:
    for p in cell.paragraphs:
        yield p
    for t in cell.tables:
        yield from _iter_table_paragraphs(t)


def _iter_block_items(container: DocxDocument | _Cell) -> Iterable[Paragraph | Table]:
    if isinstance(container, DocxDocument):
        parent_elm = container.element.body  # type: ignore[union-attr]
    elif isinstance(container, _Cell):
        parent_elm = container._tc
    else:
        parent_elm = container._element  # type: ignore[attr-defined]
    for child in parent_elm.iterchildren():
        if isinstance(child, CT_P):
            yield Paragraph(child, container)  # type: ignore[arg-type]
        elif isinstance(child, CT_Tbl):
            yield Table(child, container)  # type: ignore[arg-type]


def _para_with_inline_footnotes(p: Paragraph, footnotes: Dict[int, str]) -> str:
    parts: List[str] = []
    for run in p.runs:
        r = run._r
        refs = list(r.iter(qn("w:footnoteReference")))
        if refs:
            if run.text:
                parts.append(run.text)
            for ref in refs:
                fid = int(ref.get(qn("w:id")))
                ftxt = footnotes.get(fid, "").strip()
                parts.append(f" {ftxt} ")
            continue
        if run.text:
            parts.append(run.text)
    return "".join(parts).strip()

def _normalize(text: str) -> str:
    """Normalize text by removing artifacts and standardizing formatting.

    Args:
        text: Raw extracted text.

    Returns:
        Normalized text with consistent line breaks and punctuation.
    """
    text = text.replace("\r\n", "\n").replace("\r", "\n")
    text = _SMART_QUOTES_RE.sub('"', text)
    text = _SMART_APOSTROPHES_RE.sub("'", text)
    text = _HYPHEN_WRAP_RE.sub(r"\1\2", text)
    text = _EXCESS_BREAKS_RE.sub("\n\n", text)
    return text.strip()


def _normalize_superscripts(text: str) -> str:
    return text.translate(_SUPERSCRIPT_TRANSLATION)


def _primary_font_size(text_dict: Dict[str, Any]) -> float:
    sizes: Counter[float] = Counter()
    for block in text_dict.get("blocks", []):
        if block.get("type") != 0:
            continue
        for line in block.get("lines", []):
            for span in line.get("spans", []):
                raw_text = (span.get("text") or "").strip()
                if not raw_text:
                    continue
                size = float(span.get("size", 0.0) or 0.0)
                if size > 0:
                    sizes[round(size, 1)] += 1
    if not sizes:
        return 0.0
    return max(sizes.items(), key=lambda item: item[1])[0]


def _line_text_from_spans(spans: Iterable[Dict[str, Any]]) -> str:
    return "".join(str(span.get("text") or "") for span in spans)


def _normalize_footnote_token(token: str) -> str:
    if not token:
        return token
    start = 0
    end = len(token)
    while start < end and not token[start].isalpha():
        start += 1
    while end > start and not token[end - 1].isalpha():
        end -= 1
    if start >= end:
        return token
    leading = token[:start]
    core = token[start:end]
    trailing = token[end:]
    if not core.isupper():
        return token
    if len(core) < 2:
        return token
    if not core.isalpha():
        return token
    if len(core) < 4 and not trailing.startswith('.'):
        return token
    converted = core[0] + core[1:].lower()
    return f"{leading}{converted}{trailing}"


def _normalize_footnote_case(text: str) -> str:
    if not text:
        return text
    parts = re.split(r"(\s+)", text)
    normalized_parts: List[str] = []
    for part in parts:
        if not part:
            continue
        if part.isspace():
            normalized_parts.append(part)
        else:
            normalized_parts.append(_normalize_footnote_token(part))
    return "".join(normalized_parts)


def _span_bbox(span: Dict[str, Any]) -> Sequence[float]:
    bbox = span.get("bbox")
    if isinstance(bbox, Sequence) and len(bbox) >= 4:
        return bbox
    return (0.0, 0.0, 0.0, 0.0)


def _block_top(block: Dict[str, Any]) -> float:
    for line in block.get("lines", []):
        for span in line.get("spans", []):
            bbox = _span_bbox(span)
            return float(bbox[1])
    return float("inf")


def _block_average_font_size(block: Dict[str, Any]) -> float:
    sizes: List[float] = []
    for line in block.get("lines", []):
        for span in line.get("spans", []):
            size_val = float(span.get("size", 0.0) or 0.0)
            if size_val > 0:
                sizes.append(size_val)
    if not sizes:
        return 0.0
    return sum(sizes) / len(sizes)


def _is_footnote_block(
    block: Dict[str, Any],
    page_height: float,
    primary_font_size: float,
) -> bool:
    lines = block.get("lines", [])
    if not lines:
        return False
    normalized_lines = [
        _normalize_superscripts(_line_text_from_spans(line.get("spans", []))).strip()
        for line in lines
    ]
    footnote_starts = [line for line in normalized_lines if _FOOTNOTE_LINE_RE.match(line)]
    if not footnote_starts:
        return False
    block_top = _block_top(block)
    block_avg_size = _block_average_font_size(block)
    majority_threshold = max(1, len(lines) // 2)
    if page_height > 0 and block_top > page_height * 0.7:
        return True
    if primary_font_size > 0 and block_avg_size > 0:
        if block_avg_size <= primary_font_size * 0.85 and len(footnote_starts) >= majority_threshold:
            return True
    return False


def _parse_footnote_lines(lines: Iterable[str]) -> Dict[int, str]:
    footnotes: Dict[int, str] = {}
    current_number: Optional[int] = None
    buffer: List[str] = []

    def flush() -> None:
        nonlocal buffer
        nonlocal current_number
        if current_number is None:
            buffer = []
            return
        text = " ".join(part for part in buffer if part).strip()
        if text:
            footnotes[current_number] = _normalize_footnote_case(text)
        buffer = []

    for raw_line in lines:
        normalized = _normalize_superscripts(raw_line)
        stripped = normalized.strip()
        if not stripped:
            if buffer:
                buffer.append("")
            continue
        match = _FOOTNOTE_LINE_RE.match(stripped)
        if match:
            flush()
            number_str = match.group(1)
            digits = number_str if number_str.isdigit() else _normalize_superscripts(number_str)
            try:
                current_number = int(digits)
            except (TypeError, ValueError):
                current_number = None
            remainder = match.group(2).strip()
            buffer = [remainder] if remainder else []
        elif current_number is not None:
            buffer.append(stripped)

    if buffer:
        flush()

    return footnotes


def _ends_with_whitespace(parts: List[str]) -> bool:
    for part in reversed(parts):
        if not part:
            continue
        return part[-1].isspace()
    return False


def _render_span_with_inline_footnotes(
    span: Dict[str, Any],
    footnotes: Dict[int, str],
    primary_font_size: float,
    used: Set[int],
) -> str:
    text = span.get("text") or ""
    if not text:
        return ""
    normalized = _normalize_superscripts(text)
    matches = list(re.finditer(r"(?<!\d)(\d{1,3})(?!\d)", normalized))
    if not matches:
        return text
    font_size = float(span.get("size", 0.0) or 0.0)
    has_superscript = any(ch in _SUPERSCRIPT_CHARACTERS for ch in text)
    is_small = primary_font_size > 0 and font_size > 0 and font_size <= primary_font_size * 0.85
    if not has_superscript and not is_small:
        return text
    residual = re.sub(r"(?<!\d)\d{1,3}(?!\d)", "", normalized)
    if residual.strip():
        return text
    result: List[str] = []
    last_idx = 0
    for match in matches:
        result.append(text[last_idx:match.start()])
        num_val = int(normalized[match.start():match.end()])
        footnote_text = footnotes.get(num_val)
        if footnote_text:
            if not _ends_with_whitespace(result):
                result.append(" ")
            clean_text = footnote_text.strip()
            if clean_text:
                result.append(clean_text)
                result.append(" ")
                used.add(num_val)
        else:
            result.append(text[match.start():match.end()])
        last_idx = match.end()
    result.append(text[last_idx:])
    return "".join(result)


def _render_line_with_inline_footnotes(
    line: Dict[str, Any],
    footnotes: Dict[int, str],
    primary_font_size: float,
    used: Set[int],
) -> str:
    parts: List[str] = []
    for span in line.get("spans", []):
        parts.append(
            _render_span_with_inline_footnotes(span, footnotes, primary_font_size, used)
        )
    line_text = "".join(parts)
    line_text = re.sub(r" {2,}", " ", line_text)
    return line_text.rstrip()


def _extract_pdf_page_text(page: fitz.Page) -> str:
    try:
        text_dict: Any = page.get_text("dict")
    except Exception:  # pragma: no cover - defensive
        return page.get_text("text")
    if not isinstance(text_dict, dict):
        return page.get_text("text")

    primary_font_size = _primary_font_size(text_dict)
    page_height = float(page.rect.height)

    main_blocks: List[List[Dict[str, Any]]] = []
    footnote_lines: List[str] = []

    for block in text_dict.get("blocks", []):
        if block.get("type") != 0:
            continue
        lines = block.get("lines", [])
        if not lines:
            continue
        if _is_footnote_block(block, page_height, primary_font_size):
            for line in lines:
                footnote_lines.append(_line_text_from_spans(line.get("spans", [])))
        else:
            main_blocks.append(lines)

    footnotes = _parse_footnote_lines(footnote_lines)
    used_footnotes: Set[int] = set()

    if not main_blocks:
        if footnotes:
            logger.debug(
                "Detected footnotes without main text on page %s", getattr(page, "number", 0) + 1
            )
        return page.get_text("text")

    lines_out: List[str] = []
    for block_lines in main_blocks:
        block_texts: List[str] = []
        for line in block_lines:
            block_texts.append(
                _render_line_with_inline_footnotes(line, footnotes, primary_font_size, used_footnotes)
            )
        if block_texts:
            if lines_out and lines_out[-1] != "":
                lines_out.append("")
            lines_out.extend(block_texts)

    unused = [num for num in sorted(footnotes) if num not in used_footnotes and footnotes[num]]
    if unused:
        if lines_out:
            lines_out.append("")
        for num in unused:
            lines_out.append(footnotes[num])

    return "\n".join(lines_out)


def extract_pdf_text(file: FileStorage) -> str:
    """Extract text from PDF file, inserting footnotes inline when present."""

    file.stream.seek(0)
    page_texts: List[str] = []

    try:
        pdf_bytes = file.stream.read()
        with fitz.open(stream=pdf_bytes, filetype="pdf") as doc:
            for page in doc:
                page_text = _extract_pdf_page_text(page)
                if not page_text.strip():
                    raw_text = page.get_text("text")
                    if raw_text.strip():
                        page_text = raw_text
                    else:
                        pix = page.get_pixmap()
                        img = Image.frombytes(
                            mode="RGB",
                            size=(pix.width, pix.height),
                            data=pix.samples,
                        )
                        page_text = pytesseract.image_to_string(img)
                page_texts.append(page_text.strip())
    except Exception as exc:  # pragma: no cover - pass through for callers
        raise ValueError(f"Failed to extract text from PDF: {exc}") from exc

    combined = "\n\n\f\n\n".join(page_texts).strip()
    return _normalize(combined)

def _extract_docx_with_footnotes(doc: DocxDocument) -> str:
    """Return DOCX body text with footnotes inserted inline at their references.

    Footnote content is inserted at each reference as: "[n: footnote text]".
    Footnote numbering follows the order of first appearance in the document.

    Args:
        path: Filesystem path to a .docx file.

    Returns:
        A single string containing paragraph text with inline footnotes.

    Raises:
        FileNotFoundError: If the file does not exist.
        ValueError: If the file does not have a .docx extension.
    """
    footnotes = _load_footnotes_map(doc)


    lines: List[str] = []

    for block in _iter_block_items(doc):
        if isinstance(block, Paragraph):
            t = _para_with_inline_footnotes(block, footnotes)
            if t:
                lines.append(t)
        else:
            for par in _iter_table_paragraphs(block):
                t = _para_with_inline_footnotes(par, footnotes)
                if t:
                    lines.append(t)

    return "\n\n".join(lines)



def extract_docx_text(file: FileStorage) -> str:
    """Extract text from DOCX file, including footnotes inline.

    Args:
        file: FileStorage object containing DOCX data.

    Returns:
        Extracted and normalized text with footnotes inline.

    Raises:
        ValueError: If DOCX cannot be opened or processed.
    """
    file.stream.seek(0)

    try:
        doc = Document(file.stream)
    except Exception as exc:
        raise ValueError(f"Failed to open DOCX file: {exc}") from exc

    full_text = _extract_docx_with_footnotes(doc)
    return _normalize(full_text)


def extract_text(file: FileStorage) -> str:
    """Extract text from uploaded file based on file extension.

    Supports PDF, DOCX, and TXT files. Includes footnote extraction
    for PDF and DOCX formats.

    Args:
        file: FileStorage object containing the uploaded file.

    Returns:
        Extracted and normalized text content.

    Raises:
        ValueError: If file format is unsupported or extraction fails.
    """
    filename = file.filename or ""
    file.stream.seek(0)
    _, ext = os.path.splitext(filename.lower())

    if ext == ".txt":
        try:
            raw_text = file.stream.read().decode("utf-8")
            return _normalize(raw_text)
        except Exception as exc:
            raise ValueError(f"Failed to read TXT file: {exc}") from exc
    elif ext == ".pdf":
        return extract_pdf_text(file)
    elif ext == ".docx":
        return extract_docx_text(file)
    else:
        raise ValueError(
            f"Unsupported file format: {ext}. Supported formats: .pdf, .docx, .txt"
        )
